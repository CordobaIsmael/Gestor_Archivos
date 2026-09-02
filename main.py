from fastapi import FastAPI, Depends, HTTPException, Query
from sqlalchemy.orm import Session
from typing import List, Optional
from datetime import date
from pydantic import BaseModel
import hashlib

from services.db import get_db, init_db
from models.database import Reparto, Usuario, Caja
from core.organizer import process_incoming_folders, resolve_revision_folder

app = FastAPI(
    title="GestorArchivo API",
    description="Backend API for processing, organizing, and manually resolving file transfers.",
    version="1.0.0"
)

# Startup event to initialize tables
@app.on_event("startup")
def on_startup():
    init_db()

# ----------------- AUTH SCHEMAS & ENDPOINTS -----------------
class LoginRequest(BaseModel):
    legajo: str
    password: str

class CreateUserRequest(BaseModel):
    legajo: str
    nombre: str
    password: str
    rol: Optional[str] = "OPERADOR"

class ResetPasswordRequest(BaseModel):
    password_nueva: str

class ChangePasswordRequest(BaseModel):
    password_actual: str
    password_nueva: str

@app.post("/api/auth/login", summary="Login with Legajo and Password/DNI")
def login(data: LoginRequest, db: Session = Depends(get_db)):
    legajo_clean = data.legajo.strip()
    user = db.query(Usuario).filter(Usuario.legajo == legajo_clean).first()
    if not user:
        raise HTTPException(status_code=401, detail="Usuario / Legajo no encontrado.")
        
    if not user.activo:
        raise HTTPException(status_code=403, detail="El usuario se encuentra desactivado. Contacta al administrador.")
        
    pwd_hash = hashlib.sha256(data.password.strip().encode('utf-8')).hexdigest()
    if user.password_hash != pwd_hash:
        raise HTTPException(status_code=401, detail="Contraseña / DNI incorrecto.")
        
    return {
        "status": "success",
        "user": user.to_dict()
    }

@app.get("/api/auth/users", summary="List all operators/users (Admin only)")
def list_users(db: Session = Depends(get_db)):
    users = db.query(Usuario).order_by(Usuario.legajo.asc()).all()
    user_list = []
    for u in users:
        u_dict = u.to_dict()
        # Count repartos and closed boxes
        u_dict["total_repartos"] = db.query(Reparto).filter(Reparto.usuario_id == u.id).count()
        u_dict["total_cajas"] = db.query(Caja).filter(Caja.usuario_id == u.id).count()
        user_list.append(u_dict)
    return user_list

@app.post("/api/auth/users/create", summary="Create new operator user")
def create_user(data: CreateUserRequest, db: Session = Depends(get_db)):
    legajo_clean = data.legajo.strip()
    existing = db.query(Usuario).filter(Usuario.legajo == legajo_clean).first()
    if existing:
        raise HTTPException(status_code=400, detail=f"Ya existe un usuario con el Legajo '{legajo_clean}'.")
        
    pwd_hash = hashlib.sha256(data.password.strip().encode('utf-8')).hexdigest()
    new_user = Usuario(
        legajo=legajo_clean,
        nombre=data.nombre.strip(),
        password_hash=pwd_hash,
        rol=data.rol.upper() if data.rol else "OPERADOR",
        activo=True
    )
    db.add(new_user)
    db.commit()
    db.refresh(new_user)
    return {"status": "success", "user": new_user.to_dict()}

@app.post("/api/auth/users/{user_id}/toggle", summary="Toggle operator active status")
def toggle_user_active(user_id: int, db: Session = Depends(get_db)):
    user = db.query(Usuario).filter(Usuario.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="Usuario no encontrado.")
    user.activo = not user.activo
    db.commit()
    return {"status": "success", "user": user.to_dict()}

@app.post("/api/auth/users/{user_id}/reset-password", summary="Admin reset operator password")
def reset_user_password(user_id: int, data: ResetPasswordRequest, db: Session = Depends(get_db)):
    user = db.query(Usuario).filter(Usuario.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="Usuario no encontrado.")
    user.password_hash = hashlib.sha256(data.password_nueva.strip().encode('utf-8')).hexdigest()
    db.commit()
    return {"status": "success", "message": f"Contraseña actualizada para el usuario {user.legajo}."}

@app.post("/api/auth/users/{user_id}/change-password", summary="Operator change own password")
def change_own_password(user_id: int, data: ChangePasswordRequest, db: Session = Depends(get_db)):
    user = db.query(Usuario).filter(Usuario.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="Usuario no encontrado.")
    current_hash = hashlib.sha256(data.password_actual.strip().encode('utf-8')).hexdigest()
    if user.password_hash != current_hash:
        raise HTTPException(status_code=400, detail="La contraseña actual es incorrecta.")
    user.password_hash = hashlib.sha256(data.password_nueva.strip().encode('utf-8')).hexdigest()
    db.commit()
    return {"status": "success", "message": "Contraseña cambiada exitosamente."}

# ----------------- REPARTOS & PROCESS SCHEMAS -----------------
class ResolveRequest(BaseModel):
    empresa: str
    fecha: date
    sucursal: str
    nro_reparto: str
    salida_path: Optional[str] = None
    resolucion_guias_faltantes: Optional[dict] = None
    resolucion_guias_sin_firma: Optional[dict] = None
    modo_historico: Optional[bool] = False
    usuario_id: Optional[int] = None
    usuario_legajo: Optional[str] = None

class ProcessRequest(BaseModel):
    path: Optional[str] = None
    salida_path: Optional[str] = None
    modo_historico: Optional[bool] = False
    usuario_id: Optional[int] = None
    usuario_legajo: Optional[str] = None

@app.post("/api/process", summary="Process incoming folders")
def process_folders(data: Optional[ProcessRequest] = None, db: Session = Depends(get_db)):
    try:
        custom_path = None
        custom_salida = None
        modo_historico = False
        usuario_id = None
        usuario_legajo = None
        
        if data:
            from pathlib import Path
            if data.path:
                p = Path(data.path)
                if not p.exists() or not p.is_dir():
                    raise HTTPException(
                        status_code=400, 
                        detail=f"La ruta especificada no existe o no es un directorio: {data.path}"
                    )
                custom_path = p
                
            if data.salida_path:
                custom_salida = Path(data.salida_path)
            
            modo_historico = data.modo_historico or False
            usuario_id = data.usuario_id
            usuario_legajo = data.usuario_legajo
            
        results = process_incoming_folders(
            db, 
            custom_path=custom_path, 
            custom_salida=custom_salida,
            modo_historico=modo_historico,
            usuario_id=usuario_id,
            usuario_legajo=usuario_legajo
        )
        return {
            "status": "success",
            "message": "Procesamiento completado.",
            "data": results
        }
    except HTTPException as he:
        raise he
    except ValueError as ve:
        raise HTTPException(status_code=400, detail=str(ve))
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/repartos", summary="Get list of repartos")
def get_repartos(
    estado: Optional[str] = Query(None, description="Filter by status: ORGANIZADO, EN_REVISION"),
    usuario_id: Optional[int] = Query(None, description="Filter by operator user ID"),
    db: Session = Depends(get_db)
):
    query = db.query(Reparto)
    if estado:
        query = query.filter(Reparto.estado == estado.upper())
    if usuario_id:
        query = query.filter(Reparto.usuario_id == usuario_id)
    
    repartos = query.order_by(Reparto.fecha_procesamiento.desc()).all()
    return [r.to_dict() for r in repartos]

@app.post("/api/repartos/{reparto_id}/resolve", summary="Resolve a folder in REVISION status")
def resolve_reparto(
    reparto_id: int, 
    data: ResolveRequest, 
    db: Session = Depends(get_db)
):
    try:
        from pathlib import Path
        custom_salida = Path(data.salida_path) if data.salida_path else None
        
        updated_reparto = resolve_revision_folder(
            reparto_id=reparto_id,
            empresa=data.empresa,
            fecha_obj=data.fecha,
            sucursal=data.sucursal,
            nro_reparto=data.nro_reparto,
            db=db,
            custom_salida=custom_salida,
            resolucion_guias_faltantes=data.resolucion_guias_faltantes,
            resolucion_guias_sin_firma=data.resolucion_guias_sin_firma,
            modo_historico=data.modo_historico or False,
            usuario_id=data.usuario_id,
            usuario_legajo=data.usuario_legajo
        )
        return {
            "status": "success",
            "message": f"Reparto #{reparto_id} resolved successfully.",
            "data": updated_reparto
        }
    except ValueError as ve:
        raise HTTPException(status_code=400, detail=str(ve))
    except FileNotFoundError as fnfe:
        raise HTTPException(status_code=404, detail=str(fnfe))
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/repartos/{reparto_id}/open", summary="Open folder in File Explorer")
def open_reparto_folder(reparto_id: int, db: Session = Depends(get_db)):
    try:
        reparto = db.query(Reparto).filter(Reparto.id == reparto_id).first()
        if not reparto:
            raise HTTPException(status_code=404, detail="Reparto no encontrado.")
            
        path_to_open = reparto.ruta_nueva if reparto.ruta_nueva else reparto.ruta_original
        if not path_to_open:
            raise HTTPException(status_code=400, detail="La carpeta no tiene una ruta válida asignada.")
            
        from pathlib import Path
        import os
        
        p = Path(path_to_open)
        if not p.exists():
            raise HTTPException(
                status_code=404, 
                detail=f"La carpeta no existe físicamente en el sistema: {path_to_open}"
            )
            
        # Open folder in Windows File Explorer
        os.startfile(str(p.resolve()))
        return {
            "status": "success",
            "message": f"Carpeta abierta en el explorador: {path_to_open}"
        }
    except HTTPException as he:
        raise he
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"No se pudo abrir la carpeta: {str(e)}")

# ----------------- BOXES (CAJAS) SCHEMAS & ENDPOINTS -----------------
class NewCajaRequest(BaseModel):
    codigo: Optional[str] = None
    usuario_id: Optional[int] = None
    usuario_legajo: Optional[str] = None

class CloseCajaRequest(BaseModel):
    usuario_id: Optional[int] = None

@app.get("/api/cajas/active", summary="Get active box for specific user or global")
def get_active_caja_endpoint(usuario_id: Optional[int] = Query(None), db: Session = Depends(get_db)):
    query = db.query(Caja).filter(Caja.estado == "ACTIVA")
    if usuario_id:
        caja = query.filter(Caja.usuario_id == usuario_id).first()
        if not caja:
            # Check for generic unassigned active box
            caja = query.filter(Caja.usuario_id.is_(None)).first()
    else:
        caja = query.first()
        
    return caja.to_dict() if caja else None

@app.post("/api/cajas/new", summary="Open a new active box for an operator")
def open_new_caja(data: Optional[NewCajaRequest] = None, db: Session = Depends(get_db)):
    import datetime
    import re
    
    usuario_id = data.usuario_id if data else None
    usuario_legajo = data.usuario_legajo if data else None
    
    # Check if there is already an active box for this user
    if usuario_id:
        active_caja = db.query(Caja).filter(Caja.estado == "ACTIVA", Caja.usuario_id == usuario_id).first()
    else:
        active_caja = db.query(Caja).filter(Caja.estado == "ACTIVA").first()
        
    if active_caja:
        # Auto-close it
        active_caja.estado = "CERRADA"
        active_caja.fecha_cierre = datetime.datetime.utcnow()
        
    # Determine the code
    codigo = None
    if data and data.codigo:
        codigo = data.codigo.upper().strip()
        
    if not codigo:
        # Autogenerate the next consecutive code based on the last box in the database
        last_caja = db.query(Caja).order_by(Caja.id.desc()).first()
        next_num = 1
        if last_caja:
            match = re.search(r'(\d+)', last_caja.codigo)
            if match:
                try:
                    next_num = int(match.group(1)) + 1
                except ValueError:
                    pass
            else:
                next_num = db.query(Caja).count() + 1
        codigo = f"CAJA-{next_num:03d}"
        
    # Create new box
    new_caja = Caja(
        codigo=codigo,
        estado="ACTIVA",
        usuario_id=usuario_id,
        usuario_legajo=usuario_legajo
    )
    try:
        db.add(new_caja)
        db.commit()
        db.refresh(new_caja)
        return {"status": "success", "data": new_caja.to_dict()}
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=400, detail=f"No se pudo crear la caja (puede que el código '{codigo}' ya exista): {e}")

@app.post("/api/cajas/active/close", summary="Close active box and generate Word label")
def close_active_caja(data: Optional[CloseCajaRequest] = None, db: Session = Depends(get_db)):
    import datetime
    
    usuario_id = data.usuario_id if data else None
    
    if usuario_id:
        active_caja = db.query(Caja).filter(Caja.estado == "ACTIVA", Caja.usuario_id == usuario_id).first()
        if not active_caja:
            active_caja = db.query(Caja).filter(Caja.estado == "ACTIVA").first()
    else:
        active_caja = db.query(Caja).filter(Caja.estado == "ACTIVA").first()
        
    if not active_caja:
        raise HTTPException(status_code=400, detail="No hay ninguna caja activa para cerrar.")
        
    codigo_caja = active_caja.codigo
    
    # Update database
    active_caja.estado = "CERRADA"
    active_caja.fecha_cierre = datetime.datetime.utcnow()
    db.commit()
    db.refresh(active_caja)
    
    # Generate Word document
    try:
        from pathlib import Path
        from config.settings import settings
        
        etiquetas_dir = Path(settings.SALIDA) / "Etiquetas"
        etiquetas_dir.mkdir(parents=True, exist_ok=True)
        
        docx_path = etiquetas_dir / f"ETIQUETA_{codigo_caja}.docx"
        
        # Build Word file
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Page margins: 1 inch
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            
        # Document title
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_before = Pt(50)
        
        run_title = p_title.add_run("CAJA DE ARCHIVO FISICO")
        run_title.font.name = 'Arial'
        run_title.font.size = Pt(28)
        run_title.bold = True
        
        # Box Code (Huge Font)
        p_code = doc.add_paragraph()
        p_code.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_code.paragraph_format.space_before = Pt(70)
        p_code.paragraph_format.space_after = Pt(70)
        
        run_code = p_code.add_run(codigo_caja)
        run_code.font.name = 'Arial'
        run_code.font.size = Pt(72)
        run_code.bold = True
        
        # Footer (Date of closure & operator)
        p_date = doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        fecha_cierre_str = active_caja.fecha_cierre.strftime("%d/%m/%Y")
        op_info = f" | Operador: {active_caja.usuario_legajo}" if active_caja.usuario_legajo else ""
        run_date = p_date.add_run(f"Fecha de Cierre: {fecha_cierre_str}{op_info}")
        run_date.font.name = 'Arial'
        run_date.font.size = Pt(16)
        run_date.italic = True
        
        # Save the document
        doc.save(str(docx_path.resolve()))
        
        # Open the document automatically in MS Word on Windows
        import os
        try:
            os.startfile(str(docx_path.resolve()))
        except Exception as start_err:
            print(f"No se pudo abrir automáticamente el archivo Word: {start_err}")
            
        return {
            "status": "success",
            "message": f"Caja {codigo_caja} cerrada con éxito.",
            "file_path": str(docx_path.resolve()),
            "data": active_caja.to_dict()
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error al generar la etiqueta de Word: {e}")
